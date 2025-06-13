# import os
# import io
# import base64
# import json
# import traceback # For detailed traceback logging
# from dotenv import load_dotenv # To load .env file

# # IMPORTS for Flask and Flask-SocketIO
# from flask import Flask, request, jsonify
# from flask_socketio import SocketIO, emit, join_room, leave_room
# from flask_cors import CORS # For handling CORS with Flask

# # Patch standard library to be non-blocking with eventlet.
# # This MUST be called as early as possible.
# import eventlet
# eventlet.monkey_patch()

# # Imports for document parsing and AI calls
# import fitz  # PyMuPDF for PDF text extraction
# from reportlab.pdfgen import canvas
# from reportlab.lib.pagesizes import letter
# from reportlab.lib.units import inch
# from reportlab.platypus import Paragraph, Spacer, Frame # Added Frame for multi-flowable support in PDF
# from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
# from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
# from reportlab.platypus import SimpleDocTemplate # Added for building PDF with flowables
# import httpx # Using httpx for LM Studio API calls
# from docx import Document # For .docx files
# import openpyxl # For .xlsx files

# # Load environment variables from .env file
# load_dotenv()

# # Initialize Flask app
# app = Flask(__name__)

# # Configure Flask-CORS (equivalent to FastAPI's CORSMiddleware)
# CORS(app, resources={r"/*": {"origins": "*"}}) # Allow all origins for development. Restrict in production.

# # Initialize Flask-SocketIO
# # async_mode can be 'threading', 'eventlet', or 'gevent'
# # 'eventlet' or 'gevent' are recommended for production for performance.
# # With sync LM Studio calls, 'eventlet' is suitable.
# sio = SocketIO(app, cors_allowed_origins="*", async_mode='eventlet', logger=True, engineio_logger=True)

# # Placeholder for API key (not used with local LM Studio, but kept for consistency)
# API_KEY = os.environ.get("GEMINI_API_KEY", "")

# # Define a constant for the literal JSON error message
# JSON_ERROR_RESPONSE_LITERAL = '{"error": "JSON output malformed."}'

# # Define the model name globally
# GLOBAL_MODEL_NAME = "gemma3:4b" # Still using gemma3:4b as per your instruction

# # LM Studio API endpoint
# LM_STUDIO_API_URL = "http://192.168.234.1:1234/v1/chat/completions"

# # File size limit (10 MB)
# MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024 # 10 MB

# # Allowed MIME types (for server-side validation)
# ALLOWED_MIME_TYPES = {
#     'application/pdf': '.pdf',
#     'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
#     'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': '.xlsx',
# }

# # --- Utility Functions for Text Extraction ---

# def extract_text_from_pdf(pdf_bytes):
#     """
#     Extracts text content from PDF bytes using PyMuPDF.
#     """
#     text_content = ""
#     try:
#         doc = fitz.open(stream=pdf_bytes, filetype="pdf")
#         for page_num in range(len(doc)):
#             page = doc.load_page(page_num)
#             text_content += page.get_text()
#         doc.close()
#     except Exception as e:
#         print(f"Error extracting text from PDF: {e}")
#         return None
#     return text_content

# def extract_text_from_docx(docx_bytes):
#     """
#     Extracts text content from DOCX bytes using python-docx.
#     """
#     text_content = ""
#     try:
#         doc = Document(io.BytesIO(docx_bytes))
#         for paragraph in doc.paragraphs:
#             text_content += paragraph.text + "\n"
#     except Exception as e:
#         print(f"Error extracting text from DOCX: {e}")
#         return None
#     return text_content

# def extract_text_from_xlsx(xlsx_bytes):
#     """
#     Extracts text content from XLSX bytes using openpyxl.
#     Iterates through all sheets and cells.
#     """
#     text_content = ""
#     try:
#         workbook = openpyxl.load_workbook(io.BytesIO(xlsx_bytes))
#         for sheet_name in workbook.sheetnames:
#             sheet = workbook[sheet_name]
#             text_content += f"--- Sheet: {sheet_name} ---\n"
#             for row in sheet.iter_rows():
#                 row_values = []
#                 for cell in row:
#                     if cell.value is not None:
#                         row_values.append(str(cell.value))
#                 if row_values:
#                     text_content += " ".join(row_values) + "\n"
#             text_content += "\n"
#     except Exception as e:
#         print(f"Error extracting text from XLSX: {e}")
#         return None
#     return text_content

# def generate_pdf_from_text_basic(text_content):
#     """
#     Generates a new PDF from a given string of text using ReportLab.
#     The content is wrapped in a Paragraph for basic text flow.
#     """
#     buffer = io.BytesIO()
#     p = canvas.Canvas(buffer, pagesize=letter)
#     styles = getSampleStyleSheet()
#     style_normal = styles['Normal']

#     left_margin = inch
#     top_margin = letter[1] - inch
#     content_width = letter[0] - 2 * inch
#     line_height = 14

#     paragraphs = text_content.split('\n')
#     current_y = top_margin

#     for para_text in paragraphs:
#         paragraph = Paragraph(para_text, style_normal)
#         text_lines = paragraph.wrapOn(p, content_width, 0)[1] / line_height
#         paragraph_height = text_lines * line_height

#         if current_y - paragraph_height < inch:
#             p.showPage()
#             current_y = top_margin

#         paragraph.drawOn(p, left_margin, current_y - paragraph_height)
#         current_y -= paragraph_height + 6

#     p.save()
#     buffer.seek(0)
#     return buffer.getvalue()


# def apply_edits_to_template(edited_text: str, revised_summary: str, template_name: str) -> bytes:
#     """
#     Applies edited resume text and revised summary to a selected template (conceptually)
#     and generates a PDF. This function uses ReportLab for basic PDF generation.

#     Template options (`template_name`) are just for demonstration here, affecting
#     the PDF's title and basic styling, not loading actual visual templates.
#     """
#     buffer = io.BytesIO()
#     doc = SimpleDocTemplate(buffer, pagesize=letter,
#                             rightMargin=inch, leftMargin=inch,
#                             topMargin=inch, bottomMargin=inch)
    
#     styles = getSampleStyleSheet()

#     style_heading = ParagraphStyle(
#         'Heading',
#         parent=styles['h1'],
#         fontSize=24,
#         leading=28,
#         alignment=TA_CENTER,
#         spaceAfter=14
#     )
#     style_summary = ParagraphStyle(
#         'Summary',
#         parent=styles['Normal'],
#         fontSize=11,
#         leading=13,
#         alignment=TA_CENTER,
#         spaceAfter=12,
#         fontName='Helvetica-Oblique'
#     )
#     style_normal = ParagraphStyle(
#         'Normal',
#         parent=styles['Normal'],
#         fontSize=10,
#         leading=12,
#         alignment=TA_JUSTIFY,
#         spaceAfter=6
#     )
#     style_bullet = ParagraphStyle(
#         'Bullet',
#         parent=styles['Normal'],
#         fontSize=10,
#         leading=12,
#         alignment=TA_LEFT,
#         leftIndent=0.5 * inch,
#         bulletIndent=0.25 * inch,
#         bulletFontName='Helvetica',
#         spaceAfter=3
#     )
#     style_h2 = ParagraphStyle(
#         'H2',
#         parent=styles['h2'],
#         fontSize=14,
#         leading=16,
#         alignment=TA_LEFT,
#         spaceAfter=10
#     )

#     story = []

#     story.append(Paragraph(f"AI-Enhanced Resume ({template_name})", style_heading))
    
#     if revised_summary:
#         story.append(Paragraph("Professional Summary", style_h2))
#         story.append(Paragraph(revised_summary, style_summary))
#         story.append(Spacer(1, 0.2 * inch))

#     story.append(Paragraph("--- Main Resume Content ---", style_h2))

#     lines = edited_text.split('\n')
#     for line in lines:
#         stripped_line = line.strip()
#         if stripped_line.startswith('- '):
#             story.append(Paragraph(stripped_line.lstrip('- ').strip(), style_bullet))
#         elif stripped_line:
#             story.append(Paragraph(stripped_line, style_normal))
#         else:
#             story.append(Spacer(1, 0.1 * inch))

#     try:
#         doc.build(story)
#     except Exception as e:
#         print(f"Error building PDF with ReportLab: {e}")
#         traceback.print_exc()
#         fallback_buffer = io.BytesIO()
#         p_fallback = canvas.Canvas(fallback_buffer, pagesize=letter)
#         p_fallback.drawString(inch, letter[1] - inch, "Error generating templated PDF.")
#         p_fallback.drawString(inch, letter[1] - inch - 20, "Please see backend logs for details.")
#         p_fallback.drawString(inch, letter[1] - inch - 40, "Generated Content (Basic):")
#         p_fallback.drawString(inch, letter[1] - inch - 60, edited_text[:500] + "...")
#         p_fallback.save()
#         fallback_buffer.seek(0)
#         return fallback_buffer.getvalue()

#     buffer.seek(0)
#     return buffer.getvalue()


# def call_lm_studio_api(prompt: str, response_schema: dict = None):
#     """
#     Makes a synchronous call to LM Studio's OpenAI-compatible API for the specified model.
#     Handles cases where the model might output multiple concatenated JSON objects.
#     """
#     model_name = GLOBAL_MODEL_NAME
#     raw_response_content = None
    
#     try:
#         messages = [{"role": "user", "content": prompt}]

#         payload = {
#             "model": model_name,
#             "messages": messages,
#             "temperature": 0.7,
#             "top_p": 0.7,
#             "max_tokens": 4096,
#         }

#         # httpx client is now cooperative thanks to eventlet.monkey_patch()
#         with httpx.Client(timeout=300.0) as client:
#             response = client.post(LM_STUDIO_API_URL, json=payload)
#             response.raise_for_status()

#             full_lm_studio_response = response.json()
#             raw_response_content = full_lm_studio_response.get('choices', [{}])[0].get('message', {}).get('content', '')

#         if not raw_response_content:
#             print(f"ERROR: LM Studio model {model_name} returned empty content. Full LM Studio response: {full_lm_studio_response}")
#             return {"error": f"LM Studio model {model_name} returned an empty response. This might indicate a problem with the model or insufficient resources.", "raw_response": str(full_lm_studio_response)}
        
#         if raw_response_content.strip() == JSON_ERROR_RESPONSE_LITERAL:
#             print(f"ERROR: LM STUDIO model {model_name} explicitly returned JSON_ERROR_RESPONSE_LITERAL.")
#             return {"error": "AI Model explicitly indicated malformed JSON output.", "raw_response": raw_response_content}

#         result_text = raw_response_content
#         if result_text.startswith('```json') and result_text.endswith('```'):
#             result_text = result_text.removeprefix('```json').removesuffix('```').strip()
#         elif result_text.startswith('```') and result_text.endswith('```'):
#             result_text = result_text.removeprefix('```').removesuffix('```').strip()

#         if response_schema:
#             final_result = {}
#             decoder = json.JSONDecoder()
#             idx = 0
#             while idx < len(result_text):
#                 try:
#                     obj, new_idx = decoder.raw_decode(result_text[idx:])
#                     if isinstance(obj, dict):
#                         final_result.update(obj)
#                     idx += new_idx
#                     while idx < len(result_text) and result_text[idx].isspace():
#                         idx += 1
#                 except json.JSONDecodeError as e:
#                     print(f"DEBUG: JSONDecodeError at index {idx}. Error: {e}. Raw text segment causing error: '{result_text[idx:idx+50]}...'")
#                     next_brace = result_text.find('{', idx)
#                     if next_brace != -1:
#                         idx = next_brace
#                     else:
#                         break
#                 except Exception as e:
#                     print(f"DEBUG: Unexpected error during JSON decoding loop: {e}. Raw text: {result_text}")
#                     traceback.print_exc()
#                     break 

#             if not final_result:
#                 print(f"WARNING: Expected JSON from {model_name}, but failed to parse and merge any valid JSON objects. Raw response after cleaning: '{result_text}'")
#                 return {"error": f"Failed to parse JSON from {model_name} API. Ensure {model_name} is outputting valid JSON or can be merged.", "raw_response": result_text}
            
#             return final_result
#         else:
#             return result_text

#     except httpx.TimeoutException:
#         error_message = f"LM Studio API call timed out after 300 seconds for model {model_name}. The model might be taking too long to generate a response or is stuck."
#         print(f"ERROR: {error_message}")
#         traceback.print_exc()
#         return {"error": error_message, "raw_response": "Timeout"}
#     except httpx.HTTPStatusError as e:
#         error_message = f"HTTP error from LM Studio API for model {model_name}: {e.response.status_code} - {e.response.text}"
#         print(f"ERROR: {error_message}")
#         traceback.print_exc()
#         return {"error": error_message, "raw_response": e.response.text}
#     except httpx.RequestError as e:
#         error_message = f"Network or request error connecting to LM Studio API for model {model_name}: {e}"
#         print(f"ERROR: {e}")
#         traceback.print_exc()
#         return {"error": error_message, "raw_response": str(e)}
#     except Exception as e:
#         error_message = f"General Exception during LM STUDIO API call for {model_name}: {e}"
#         print(f"ERROR: {e}")
#         traceback.print_exc()
#         return {"error": error_message, "raw_response": str(e)}
#     finally:
#         print(f"INFO: Finished call_lm_studio_api for {model_name}.")
#         if raw_response_content:
#             print(f"INFO: Full raw LM Studio content (first 200 chars): {str(raw_response_content)[:200]}...")
#         else:
#             print("INFO: No raw LM Studio content captured or content was empty.")


# def async_lm_studio_call(sid, prompt, schema, event_name, original_resume_text=None):
#     """
#     Handles the LM Studio API call and emits the result to the client.
#     This function is run as a background task via sio.start_background_task,
#     and because eventlet.monkey_patch() is used, blocking I/O within it
#     will be cooperative.
#     """
#     with app.app_context(): # Ensure Flask app context is available for emit
#         try:
#             analysis_result = call_lm_studio_api(prompt, schema) # Direct call, now cooperative

#             if "error" in analysis_result:
#                 sio.emit('error', {'message': f"AI Analysis Error: {analysis_result['error']}"}, room=sid, namespace='/')
#                 return

#             if event_name == 'analysis_result':
#                 sio.emit(event_name, {
#                     'score': analysis_result.get('score', 0),
#                     'suggestions': analysis_result.get('suggestions', []),
#                     'revised_summary': analysis_result.get('revised_summary', ''),
#                     'extracted_resume_text': original_resume_text, # Pass original text
#                     'ai_revised_full_resume_text': analysis_result.get('ai_revised_full_resume_text', '')
#                 }, room=sid, namespace='/')
#             elif event_name == 'interview_prep_materials':
#                 sio.emit(event_name, analysis_result, room=sid, namespace='/')
#             elif event_name == 'interview_question':
#                 sio.emit(event_name, analysis_result, room=sid, namespace='/')
#             elif event_name == 'interview_feedback':
#                 sio.emit(event_name, analysis_result, room=sid, namespace='/')
#             elif event_name == 'career_roadmap':
#                 sio.emit(event_name, analysis_result, room=sid, namespace='/')

#         except Exception as e:
#             print(f"ERROR: Exception in async_lm_studio_call: {e}")
#             traceback.print_exc()
#             sio.emit('error', {'message': f'Server error processing LLM result: {e}'}, room=sid, namespace='/')


# # --- Socket.IO Event Handlers (using Flask-SocketIO decorators) ---

# @sio.on('connect')
# def connect():
#     print(f'Client connected: {request.sid}')
#     emit('status', {'message': 'Connected to backend!'}, room=request.sid)

# @sio.on('disconnect')
# def disconnect():
#     print(f'Client disconnected: {request.sid}')

# @sio.on('upload_resume_and_jd')
# def handle_upload_resume_and_jd(data):
#     """
#     Receives resume PDF and job description, then performs analysis.
#     Handles multiple file types and size validation.
#     """
#     print(f"Received upload_resume_and_jd event from {request.sid}")
#     resume_file_b64 = data.get('resume_file')
#     job_description = data.get('job_description')
#     file_type = data.get('file_type')

#     if not resume_file_b64 or not job_description or not file_type:
#         emit('error', {'message': 'Resume file, Job Description, and file type are required.'}, room=request.sid)
#         return

#     try:
#         resume_bytes = base64.b64decode(resume_file_b64)

#         if len(resume_bytes) > MAX_FILE_SIZE_BYTES:
#             emit('error', {'message': f'File size exceeds the limit of {MAX_FILE_SIZE_BYTES / (1024 * 1024):.0f} MB.'}, room=request.sid)
#             return

#         resume_text = None
#         if file_type == 'application/pdf':
#             resume_text = extract_text_from_pdf(resume_bytes)
#         elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
#             resume_text = extract_text_from_docx(resume_bytes)
#         elif file_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
#             resume_text = extract_text_from_xlsx(resume_bytes)
#         else:
#             emit('error', {'message': 'Unsupported file type. Please upload PDF, DOCX, or XLSX.'}, room=request.sid)
#             return

#         if not resume_text:
#             emit('error', {'message': f'Could not extract text from the provided {file_type} file. It might be empty or corrupted.'}, room=request.sid)
#             return

#         analysis_prompt = f"""
# As an expert resume analyzer, compare the following resume to the provided job description.
# Your task is to:
# 1.  Assign a compatibility score between 1 and 10 (10 being a perfect match).
# 2.  Identify key areas where the resume could be improved to better align with the job description. Provide specific, actionable suggestions.
# 3.  Suggest a brief (2-3 sentences) professional summary/objective for the resume, tailored to the job description.
# 4.  Rewrite the ENTIRE provided resume content to optimize it for the job description and improve the compatibility score. Focus on incorporating relevant keywords, emphasizing achievements pertinent to the role, and making the resume more impactful for this specific job. Maintain the original structure and content where possible, but rephrase and enhance sections as needed. Ensure the revised resume content is coherent and professional.

# **IMPORTANT:** Your entire response MUST be a SINGLE JSON object. Do NOT include any other text, explanations, or markdown formatting.
# The JSON object MUST have the following keys:
# -   `score`: An integer from 1 to 10.
# -   `suggestions`: An array of strings, each string being an actionable suggestion.
# -   `revised_summary`: A brief (2-3 sentences) suggested professional summary/objective for the resume, tailored to the job description.
# -   `ai_revised_full_resume_text`: A string containing the entire rewritten resume content, optimized for the job description.
# If you cannot generate valid JSON, provide the literal string '{JSON_ERROR_RESPONSE_LITERAL}'.

# Resume:
# ```
# {resume_text}
# ```

# Job Description:
# ```
# {job_description}
# ```
# """
#         analysis_schema = {
#             "type": "OBJECT",
#             "properties": {
#                 "score": {"type": "INTEGER", "description": "Compatibility score (1-10)"},
#                 "suggestions": {"type": "ARRAY", "items": {"type": "STRING"}},
#                 "revised_summary": {"type": "STRING", "description": "Suggested professional summary/objective"},
#                 "ai_revised_full_resume_text": {"type": "STRING", "description": "The entire rewritten resume content"}
#             },
#             "required": ["score", "suggestions", "revised_summary", "ai_revised_full_resume_text"]
#         }

#         # Submit the LLM call to the eventlet background task directly
#         sio.start_background_task(async_lm_studio_call, sid=request.sid, prompt=analysis_prompt, schema=analysis_schema, event_name='analysis_result', original_resume_text=resume_text)

#     except Exception as e:
#         print(f"ERROR: Exception during upload_resume_and_jd: {e}")
#         traceback.print_exc()
#         emit('error', {'message': f'Server error during analysis: {e}'}, room=request.sid)


# @sio.on('apply_edits')
# def handle_apply_edits(data):
#     """
#     Receives edited resume text (now AI-generated and potentially user-tweaked),
#     revised summary, and template choice, then generates a new PDF.
#     """
#     print(f"Received apply_edits event from {request.sid}")
#     edited_resume_text = data.get('edited_resume_text')
#     revised_summary = data.get('revised_summary')
#     selected_template = data.get('selected_template')

#     if not edited_resume_text:
#         emit('error', {'message': 'Edited resume text is empty.'}, room=request.sid)
#         return
#     if not selected_template:
#         emit('error', {'message': 'Please select a resume template.'}, room=request.sid)
#         return

#     try:
#         new_pdf_bytes = apply_edits_to_template(edited_resume_text, revised_summary, selected_template)
#         new_pdf_b64 = base64.b64encode(new_pdf_bytes).decode('utf-8')

#         emit('updated_resume_pdf', {'pdf_b64': new_pdf_b64}, room=request.sid)

#     except Exception as e:
#         print(f"ERROR: Exception during apply_edits: {e}")
#         traceback.print_exc()
#         emit('error', {'message': f'Server error during PDF generation: {e}'}, room=request.sid)

# @sio.on('request_interview_prep')
# def handle_request_interview_prep(data):
#     """
#     Generates interview preparation materials based on job description.
#     """
#     print(f"Received request_interview_prep event from {request.sid}")
#     job_description = data.get('job_description')

#     if not job_description:
#         emit('error', {'message': 'Job Description is required for interview prep.'}, room=request.sid)
#         return

#     try:
#         interview_prompt = f"""
# As an expert interview coach, based on the following job description, generate:
# 1.  5-7 common interview questions that an applicant for this role might face.
# 2.  For each question, provide a brief, high-level outline of an ideal answer, highlighting what aspects the interviewer is looking for.
# 3.  3 general interview tips specifically relevant to preparing for an interview for this type of role.

# **IMPORTANT:** Your entire response MUST be a SINGLE JSON object. Do NOT include any other text, explanations, or markdown formatting.
# The JSON object MUST have the following keys:
# -   `interview_questions`: An array of objects, each with `question` (string) and `ideal_answer_outline` (string).
# -   `interview_tips`: An array of strings, each string being a general interview tip.
# If you cannot generate valid JSON, provide the literal string '{JSON_ERROR_RESPONSE_LITERAL}'.

# Job Description:
# ```
# {job_description}
# ```
# """
#         interview_schema = {
#             "type": "OBJECT",
#             "properties": {
#                 "interview_questions": {
#                     "type": "ARRAY",
#                     "items": {
#                         "type": "OBJECT",
#                         "properties": {
#                             "question": {"type": "STRING"},
#                             "ideal_answer_outline": {"type": "STRING"}
#                         },
#                         "required": ["question", "ideal_answer_outline"]
#                     }
#                 },
#                 "interview_tips": {"type": "ARRAY", "items": {"type": "STRING"}}
#             },
#             "required": ["interview_questions", "interview_tips"]
#         }

#         # Submit the LLM call to the eventlet background task directly
#         sio.start_background_task(async_lm_studio_call, sid=request.sid, prompt=interview_prompt, schema=interview_schema, event_name='interview_prep_materials')

#     except Exception as e:
#         print(f"ERROR: Exception during request_interview_prep: {e}")
#         traceback.print_exc()
#         emit('error', {'message': f'Server error during interview prep: {e}'}, room=request.sid)

# @sio.on('start_interview_simulation')
# def handle_start_interview_simulation(data):
#     """
#     Generates a personalized behavioral interview question based on resume and job description.
#     """
#     print(f"Received start_interview_simulation event from {request.sid}")
#     resume_text = data.get('resume_text')
#     job_description = data.get('job_description')

#     if not resume_text or not job_description:
#         emit('error', {'message': 'Resume text and Job Description are required for interview simulation.'}, room=request.sid)
#         return

#     try:
#         simulation_prompt = f"""
# As an expert interviewer. Based on the following resume and job description, generate ONE challenging behavioral interview question.
# The question should focus on a skill or experience that is crucial for the role but might be a subtle gap or a strong point to elaborate on from the resume.
# Frame it as a STAR method question (e.g., "Tell me about a time when...").
# Also, provide a brief (1-2 sentences) context or scenario related to the question.

# **IMPORTANT:** Your entire response MUST be a SINGLE JSON object. Do NOT include any other text, explanations, or markdown formatting.
# The JSON object MUST have
# -   `question`: The behavioral interview question (string).
# -   `scenario`: A brief contextual scenario for the question (string).
# If you cannot generate valid JSON, provide the literal string '{JSON_ERROR_RESPONSE_LITERAL}'.

# Resume:
# ```
# {resume_text}
# ```

# Job Description:
# ```
# {job_description}
# ```
# """
#         simulation_schema = {
#             "type": "OBJECT",
#             "properties": {
#                 "question": {"type": "STRING"},
#                 "scenario": {"type": "STRING"}
#             },
#             "required": ["question", "scenario"]
#         }

#         # Submit the LLM call to the eventlet background task directly
#         sio.start_background_task(async_lm_studio_call, sid=request.sid, prompt=simulation_prompt, schema=simulation_schema, event_name='interview_question')

#     except Exception as e:
#         print(f"ERROR: Exception during start_interview_simulation: {e}")
#         traceback.print_exc()
#         emit('error', {'message': f'Server error during simulation: {e}'}, room=request.sid)

# @sio.on('get_interview_feedback')
# def handle_get_interview_feedback(data):
#     """
#     Provides real-time feedback on a user's interview answer.
#     """
#     print(f"Received get_interview_feedback event from {request.sid}")
#     question = data.get('question')
#     user_answer = data.get('user_answer')
#     job_description = data.get('job_description')
#     resume_text = data.get('resume_text')

#     if not question or not user_answer or not job_description or not resume_text:
#         emit('error', {'message': 'Question, user answer, job description, and resume are required for feedback.'}, room=request.sid)
#         return

#     try:
#         feedback_prompt = f"""
# As an expert interview coach. Analyze the following user's interview answer for the given question, considering the job description and their resume.
# Provide constructive feedback focusing on:
# 1.  **STAR Method Adherence:** Does the answer follow the Situation, Task, Action, Result structure? Point out missing components.
# 2.  **Relevance & Impact:** How well does the answer relate to the job description and demonstrate quantifiable impact? Suggest ways to improve relevance and add metrics.
# 3.  **Clarity & Conciseness:** Is the answer clear, easy to understand, and to the point?
# 4.  **Overall Suggestion:** A final overall tip for improving the answer.

# **IMPORTANT:** Your entire response MUST be a SINGLE JSON object. Do NOT include any other text, explanations, or markdown formatting.
# The JSON object MUST have the following keys:
# -   `star_feedback`: string
# -   `relevance_impact_feedback`: string
# -   `clarity_feedback`: string
# -   `overall_suggestion`: string
# If you cannot generate valid JSON, provide the literal string '{JSON_ERROR_RESPONSE_LITERAL}'.

# ---
# Job Description:
# ```
# {job_description}
# ```

# Resume:
# ```
# {resume_text}
# ```

# Interview Question:
# ```
# {question}
# ```

# User's Answer:
# ```
# {user_answer}
# ```
# ---
# """
#         feedback_schema = {
#             "type": "OBJECT",
#             "properties": {
#                 "star_feedback": {"type": "STRING"},
#                 "relevance_impact_feedback": {"type": "STRING"},
#                 "clarity_feedback": {"type": "STRING"},
#                 "overall_suggestion": {"type": "STRING"}
#             },
#             "required": ["star_feedback", "relevance_impact_feedback", "clarity_feedback", "overall_suggestion"]
#         }

#         # Submit the LLM call to the eventlet background task directly
#         sio.start_background_task(async_lm_studio_call, sid=request.sid, prompt=feedback_prompt, schema=feedback_schema, event_name='interview_feedback')

#     except Exception as e:
#         print(f"ERROR: Exception during get_interview_feedback: {e}")
#         traceback.print_exc()
#         emit('error', {'message': f'Server error during feedback: {e}'}, room=request.sid)

# @sio.on('request_career_roadmap')
# def handle_request_career_roadmap(data):
#     """
#     Generates a predictive career trajectory and skill pathing roadmap.
#     """
#     print(f"Received request_career_roadmap event from {request.sid}")
#     resume_text = data.get('resume_text')
#     job_description = data.get('job_description')
#     desired_roles = data.get('desired_roles')

#     if not resume_text or not job_description or not desired_roles:
#         emit('error', {'message': 'Resume text, current job description, and desired roles are required for career roadmap.'}, room=request.sid)
#         return

#     desired_roles_str = ", ".join(desired_roles)

#     try:
#         roadmap_prompt = f"""
# As an expert career coach and talent development specialist.
# Based on the provided resume and the current target job description, and the user's desired future career roles ({desired_roles_str}), generate a personalized career roadmap.

# Your roadmap should:
# 1.  **Identify Key Skill Gaps:** What specific skills or experiences are missing from the resume for the user to progress towards "{desired_roles_str}"? Categorize these (e.g., Technical, Leadership, Communication, etc.).
# 2.  **Suggest Targeted Learning:** For each identified skill gap, recommend 1-2 specific types of online courses, certifications, or learning resources (e.g., "Advanced Python course on Coursera," "PMP Certification," "Public Speaking workshop").
# 3.  **Recommend Practical Projects/Experiences:** Suggest specific types of projects or work experiences (personal or professional) that would help bridge these gaps and make the resume more compelling for the desired roles.
# 4.  **Networking/Soft Skill Advice:** Provide 2-3 actionable tips related to networking, mentorship, or developing crucial soft skills for career advancement.
# 5.  **Aspirational Summary/Next Steps:** A brief summary of how their resume could look after implementing these steps, and a concluding encouraging remark.

# **IMPORTANT:** Your entire response MUST be a SINGLE JSON object. Do NOT include any other text, explanations, or markdown formatting.
# The JSON object MUST have the following keys:
# -   `skill_gaps`: An array of objects, each with `category` (string) and `skills` (array of strings).
# -   `learning_recommendations`: An array of objects, each with `skill_area` (string) and `resources` (array of strings, e.g., "Course: AWS Certified Solutions Architect - Associate (Coursera)").
# -   `project_ideas`: An array of strings, each being a project idea.
# -   `networking_tips`: An array of strings, each being a networking/soft skill tip.
# -   `aspirational_summary`: string
# If you cannot generate valid JSON, provide the literal string '{JSON_ERROR_RESPONSE_LITERAL}'.

# ---
# Resume:
# ```
# {resume_text}
# ```

# Current Job Description:
# ```
# {job_description}
# ```

# Desired Future Roles:
# ```
# {desired_roles_str}
# ```
# ---
# """
#         roadmap_schema = {
#             "type": "OBJECT",
#             "properties": {
#                 "skill_gaps": {
#                     "type": "ARRAY",
#                     "items": {
#                         "type": "OBJECT",
#                         "properties": {
#                             "category": {"type": "STRING"},
#                             "skills": {"type": "ARRAY", "items": {"type": "STRING"}}
#                         },
#                         "required": ["category", "skills"]
#                     }
#                 },
#                 "learning_recommendations": {
#                     "type": "ARRAY",
#                     "items": {
#                         "type": "OBJECT",
#                         "properties": {
#                             "skill_area": {"type": "STRING"},
#                             "resources": {"type": "ARRAY", "items": {"type": "STRING"}}
#                         },
#                         "required": ["skill_area", "resources"]
#                     }
#                 },
#                 "project_ideas": {"type": "ARRAY", "items": {"type": "STRING"}},
#                 "networking_tips": {"type": "ARRAY", "items": {"type": "STRING"}},
#                 "aspirational_summary": {"type": "STRING"}
#             },
#             "required": ["skill_gaps", "learning_recommendations", "project_ideas", "networking_tips", "aspirational_summary"]
#         }

#         # Submit the LLM call to the eventlet background task directly
#         sio.start_background_task(async_lm_studio_call, sid=request.sid, prompt=roadmap_prompt, schema=roadmap_schema, event_name='career_roadmap')

#     except Exception as e:
#         print(f"ERROR: Exception during request_career_roadmap: {e}")
#         traceback.print_exc()
#         emit('error', {'message': f'Server error during career roadmap: {e}'}, room=request.sid)


# # --- Main entry point ---
# if __name__ == '__main__':
#     # When using Flask-SocketIO with async_mode='eventlet', you run the app via sio.run()
#     # If running with a production WSGI server like Gunicorn, you would configure Gunicorn
#     # to use the eventlet/gevent worker class and bind it to Flask-SocketIO's app.
#     # Example for Gunicorn: gunicorn --worker-class eventlet -w 1 main:app
#     print("Starting Flask-SocketIO server...")
#     try:
#         sio.run(app, host='0.0.0.0', port=5000, debug=True)
#     finally:
#         # The ThreadPoolExecutor is no longer directly managed as eventlet.monkey_patch()
#         # handles cooperative multitasking for I/O. If there were other CPU-bound tasks
#         # not related to I/O that needed true threading, the executor would be useful.
#         # For now, it's removed as it's not needed for this particular issue.
#         pass
















import os
import io
import base64
import json
import traceback # For detailed traceback logging
from dotenv import load_dotenv # To load .env file

# IMPORTS for Flask and Flask-SocketIO
from flask import Flask, request, jsonify, render_template # Added render_template
from flask_socketio import SocketIO, emit, join_room, leave_room
from flask_cors import CORS # For handling CORS with Flask

# Patch standard library to be non-blocking with eventlet.
# This MUST be called as early as possible.
import eventlet
eventlet.monkey_patch()

# Imports for document parsing and AI calls
import fitz  # PyMuPDF for PDF text extraction
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, Spacer, Frame # Added Frame for multi-flowable support in PDF
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.platypus import SimpleDocTemplate # Added for building PDF with flowables
import httpx # Using httpx for LM Studio API calls
from docx import Document # For .docx files
import openpyxl # For .xlsx files

# Load environment variables from .env file
load_dotenv()

# Initialize Flask app
app = Flask(__name__)

# Configure Flask-CORS (equivalent to FastAPI's CORSMiddleware)
CORS(app, resources={r"/*": {"origins": "*"}}) # Allow all origins for development. Restrict in production.

# Initialize Flask-SocketIO
# async_mode can be 'threading', 'eventlet', or 'gevent'
# 'eventlet' or 'gevent' are recommended for production for performance.
# With sync LM Studio calls, 'eventlet' is suitable.
sio = SocketIO(app, cors_allowed_origins="*", async_mode='eventlet', logger=True, engineio_logger=True)

# Placeholder for API key (not used with local LM Studio, but kept for consistency)
API_KEY = os.environ.get("GEMINI_API_KEY", "")

# Define a constant for the literal JSON error message
JSON_ERROR_RESPONSE_LITERAL = '{"error": "JSON output malformed."}'

# Define the model name globally
GLOBAL_MODEL_NAME = "gemma3:4b" # Still using gemma3:4b as per your instruction

# LM Studio API endpoint
LM_STUDIO_API_URL = "http://192.168.234.1:1234/v1/chat/completions"

# File size limit (10 MB)
MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024 # 10 MB

# Allowed MIME types (for server-side validation)
ALLOWED_MIME_TYPES = {
    'application/pdf': '.pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': '.xlsx',
}

# --- Utility Functions for Text Extraction ---

def extract_text_from_pdf(pdf_bytes):
    """
    Extracts text content from PDF bytes using PyMuPDF.
    """
    text_content = ""
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text_content += page.get_text()
        doc.close()
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return None
    return text_content

def extract_text_from_docx(docx_bytes):
    """
    Extracts text content from DOCX bytes using python-docx.
    """
    text_content = ""
    try:
        doc = Document(io.BytesIO(docx_bytes))
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return None
    return text_content

def extract_text_from_xlsx(xlsx_bytes):
    """
    Extracts text content from XLSX bytes using openpyxl.
    Iterates through all sheets and cells.
    """
    text_content = ""
    try:
        workbook = openpyxl.load_workbook(io.BytesIO(xlsx_bytes))
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text_content += f"--- Sheet: {sheet_name} ---\n"
            for row in sheet.iter_rows():
                row_values = []
                for cell in row:
                    if cell.value is not None:
                        row_values.append(str(cell.value))
                if row_values:
                    text_content += " ".join(row_values) + "\n"
            text_content += "\n"
    except Exception as e:
        print(f"Error extracting text from XLSX: {e}")
        return None
    return text_content

def generate_pdf_from_text_basic(text_content):
    """
    Generates a new PDF from a given string of text using ReportLab.
    The content is wrapped in a Paragraph for basic text flow.
    """
    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    style_normal = styles['Normal']

    left_margin = inch
    top_margin = letter[1] - inch
    content_width = letter[0] - 2 * inch
    line_height = 14

    paragraphs = text_content.split('\n')
    current_y = top_margin

    for para_text in paragraphs:
        paragraph = Paragraph(para_text, style_normal)
        text_lines = paragraph.wrapOn(p, content_width, 0)[1] / line_height
        paragraph_height = text_lines * line_height

        if current_y - paragraph_height < inch:
            p.showPage()
            current_y = top_margin

        paragraph.drawOn(p, left_margin, current_y - paragraph_height)
        current_y -= paragraph_height + 6

    p.save()
    buffer.seek(0)
    return buffer.getvalue()


def apply_edits_to_template(edited_text: str, revised_summary: str, template_name: str) -> bytes:
    """
    Applies edited resume text and revised summary to a selected template (conceptually)
    and generates a PDF. This function uses ReportLab for basic PDF generation.

    Template options (`template_name`) are just for demonstration here, affecting
    the PDF's title and basic styling, not loading actual visual templates.
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                            rightMargin=inch, leftMargin=inch,
                            topMargin=inch, bottomMargin=inch)
    
    styles = getSampleStyleSheet()

    style_heading = ParagraphStyle(
        'Heading',
        parent=styles['h1'],
        fontSize=24,
        leading=28,
        alignment=TA_CENTER,
        spaceAfter=14
    )
    style_summary = ParagraphStyle(
        'Summary',
        parent=styles['Normal'],
        fontSize=11,
        leading=13,
        alignment=TA_CENTER,
        spaceAfter=12,
        fontName='Helvetica-Oblique'
    )
    style_normal = ParagraphStyle(
        'Normal',
        parent=styles['Normal'],
        fontSize=10,
        leading=12,
        alignment=TA_JUSTIFY,
        spaceAfter=6
    )
    style_bullet = ParagraphStyle(
        'Bullet',
        parent=styles['Normal'],
        fontSize=10,
        leading=12,
        alignment=TA_LEFT,
        leftIndent=0.5 * inch,
        bulletIndent=0.25 * inch,
        bulletFontName='Helvetica',
        spaceAfter=3
    )
    style_h2 = ParagraphStyle(
        'H2',
        parent=styles['h2'],
        fontSize=14,
        leading=16,
        alignment=TA_LEFT,
        spaceAfter=10
    )

    story = []

    story.append(Paragraph(f"AI-Enhanced Resume ({template_name})", style_heading))
    
    if revised_summary:
        story.append(Paragraph("Professional Summary", style_h2))
        story.append(Paragraph(revised_summary, style_summary))
        story.append(Spacer(1, 0.2 * inch))

    story.append(Paragraph("--- Main Resume Content ---", style_h2))

    lines = edited_text.split('\n')
    for line in lines:
        stripped_line = line.strip()
        if stripped_line.startswith('- '):
            story.append(Paragraph(stripped_line.lstrip('- ').strip(), style_bullet))
        elif stripped_line:
            story.append(Paragraph(stripped_line, style_normal))
        else:
            story.append(Spacer(1, 0.1 * inch))

    try:
        doc.build(story)
    except Exception as e:
        print(f"Error building PDF with ReportLab: {e}")
        traceback.print_exc()
        fallback_buffer = io.BytesIO()
        p_fallback = canvas.Canvas(fallback_buffer, pagesize=letter)
        p_fallback.drawString(inch, letter[1] - inch, "Error generating templated PDF.")
        p_fallback.drawString(inch, letter[1] - inch - 20, "Please see backend logs for details.")
        p_fallback.drawString(inch, letter[1] - inch - 40, "Generated Content (Basic):")
        p_fallback.drawString(inch, letter[1] - inch - 60, edited_text[:500] + "...")
        p_fallback.save()
        fallback_buffer.seek(0)
        return fallback_buffer.getvalue()

    buffer.seek(0)
    return buffer.getvalue()


def call_lm_studio_api(prompt: str, response_schema: dict = None):
    """
    Makes a synchronous call to LM Studio's OpenAI-compatible API for the specified model.
    Handles cases where the model might output multiple concatenated JSON objects.
    """
    model_name = GLOBAL_MODEL_NAME
    raw_response_content = None
    
    try:
        messages = [{"role": "user", "content": prompt}]

        payload = {
            "model": model_name,
            "messages": messages,
            "temperature": 0.7,
            "top_p": 0.7,
            "max_tokens": 4096,
        }

        # httpx client is now cooperative thanks to eventlet.monkey_patch()
        with httpx.Client(timeout=300.0) as client:
            response = client.post(LM_STUDIO_API_URL, json=payload)
            response.raise_for_status()

            full_lm_studio_response = response.json()
            raw_response_content = full_lm_studio_response.get('choices', [{}])[0].get('message', {}).get('content', '')

        if not raw_response_content:
            print(f"ERROR: LM Studio model {model_name} returned empty content. Full LM Studio response: {full_lm_studio_response}")
            return {"error": f"LM Studio model {model_name} returned an empty response. This might indicate a problem with the model or insufficient resources.", "raw_response": str(full_lm_studio_response)}
        
        if raw_response_content.strip() == JSON_ERROR_RESPONSE_LITERAL:
            print(f"ERROR: LM STUDIO model {model_name} explicitly returned JSON_ERROR_RESPONSE_LITERAL.")
            return {"error": "AI Model explicitly indicated malformed JSON output.", "raw_response": raw_response_content}

        result_text = raw_response_content
        if result_text.startswith('```json') and result_text.endswith('```'):
            result_text = result_text.removeprefix('```json').removesuffix('```').strip()
        elif result_text.startswith('```') and result_text.endswith('```'):
            result_text = result_text.removeprefix('```').removesuffix('```').strip()

        if response_schema:
            final_result = {}
            decoder = json.JSONDecoder()
            idx = 0
            while idx < len(result_text):
                try:
                    obj, new_idx = decoder.raw_decode(result_text[idx:])
                    if isinstance(obj, dict):
                        final_result.update(obj)
                    idx += new_idx
                    while idx < len(result_text) and result_text[idx].isspace():
                        idx += 1
                except json.JSONDecodeError as e:
                    print(f"DEBUG: JSONDecodeError at index {idx}. Error: {e}. Raw text segment causing error: '{result_text[idx:idx+50]}...'")
                    next_brace = result_text.find('{', idx)
                    if next_brace != -1:
                        idx = next_brace
                    else:
                        break
                except Exception as e:
                    print(f"DEBUG: Unexpected error during JSON decoding loop: {e}. Raw text: {result_text}")
                    traceback.print_exc()
                    break 

            if not final_result:
                print(f"WARNING: Expected JSON from {model_name}, but failed to parse and merge any valid JSON objects. Raw response after cleaning: '{result_text}'")
                return {"error": f"Failed to parse JSON from {model_name} API. Ensure {model_name} is outputting valid JSON or can be merged.", "raw_response": result_text}
            
            return final_result
        else:
            return result_text

    except httpx.TimeoutException:
        error_message = f"LM Studio API call timed out after 300 seconds for model {model_name}. The model might be taking too long to generate a response or is stuck."
        print(f"ERROR: {error_message}")
        traceback.print_exc()
        return {"error": error_message, "raw_response": "Timeout"}
    except httpx.HTTPStatusError as e:
        error_message = f"HTTP error from LM Studio API for model {model_name}: {e.response.status_code} - {e.response.text}"
        print(f"ERROR: {error_message}")
        traceback.print_exc()
        return {"error": error_message, "raw_response": e.response.text}
    except httpx.RequestError as e:
        error_message = f"Network or request error connecting to LM Studio API for model {model_name}: {e}"
        print(f"ERROR: {e}")
        traceback.print_exc()
        return {"error": error_message, "raw_response": str(e)}
    except Exception as e:
        error_message = f"General Exception during LM STUDIO API call for {model_name}: {e}"
        print(f"ERROR: {e}")
        traceback.print_exc()
        return {"error": error_message, "raw_response": str(e)}
    finally:
        print(f"INFO: Finished call_lm_studio_api for {model_name}.")
        if raw_response_content:
            print(f"INFO: Full raw LM Studio content (first 200 chars): {str(raw_response_content)[:200]}...")
        else:
            print("INFO: No raw LM Studio content captured or content was empty.")


def async_lm_studio_call(sid, prompt, schema, event_name, original_resume_text=None):
    """
    Handles the LM Studio API call and emits the result to the client.
    This function is run as a background task via sio.start_background_task,
    and because eventlet.monkey_patch() is used, blocking I/O within it
    will be cooperative.
    """
    with app.app_context(): # Ensure Flask app context is available for emit
        try:
            analysis_result = call_lm_studio_api(prompt, schema) # Direct call, now cooperative

            if "error" in analysis_result:
                sio.emit('error', {'message': f"AI Analysis Error: {analysis_result['error']}"}, room=sid, namespace='/')
                return

            if event_name == 'analysis_result':
                sio.emit(event_name, {
                    'score': analysis_result.get('score', 0),
                    'suggestions': analysis_result.get('suggestions', []),
                    'revised_summary': analysis_result.get('revised_summary', ''),
                    'extracted_resume_text': original_resume_text, # Pass original text
                    'ai_revised_full_resume_text': analysis_result.get('ai_revised_full_resume_text', '')
                }, room=sid, namespace='/')
            elif event_name == 'interview_prep_materials':
                sio.emit(event_name, analysis_result, room=sid, namespace='/')
            elif event_name == 'interview_question':
                sio.emit(event_name, analysis_result, room=sid, namespace='/')
            elif event_name == 'interview_feedback':
                sio.emit(event_name, analysis_result, room=sid, namespace='/')
            elif event_name == 'career_roadmap':
                sio.emit(event_name, analysis_result, room=sid, namespace='/')

        except Exception as e:
            print(f"ERROR: Exception in async_lm_studio_call: {e}")
            traceback.print_exc()
            sio.emit('error', {'message': f'Server error processing LLM result: {e}'}, room=sid, namespace='/')


# --- Flask Routes ---
@app.route('/')
def index():
    """Serves the index.html file from the templates folder."""
    return render_template('index.html')

@app.route('/test', methods=['GET'])
def test_api():
    dummy_data = {
        "id": 1,
        "name": "Dummy Data",
        "status": "success"
    }
    return jsonify(dummy_data)


# --- Socket.IO Event Handlers (using Flask-SocketIO decorators) ---

@sio.on('connect')
def connect():
    print(f'Client connected: {request.sid}')
    emit('status', {'message': 'Connected to backend!'}, room=request.sid)

@sio.on('disconnect')
def disconnect():
    print(f'Client disconnected: {request.sid}')

@sio.on('upload_resume_and_jd')
def handle_upload_resume_and_jd(data):
    """
    Receives resume PDF and job description, then performs analysis.
    Handles multiple file types and size validation.
    """
    print(f"Received upload_resume_and_jd event from {request.sid}")
    resume_file_b64 = data.get('resume_file')
    job_description = data.get('job_description')
    file_type = data.get('file_type')

    if not resume_file_b64 or not job_description or not file_type:
        emit('error', {'message': 'Resume file, Job Description, and file type are required.'}, room=request.sid)
        return

    try:
        resume_bytes = base64.b64decode(resume_file_b64)

        if len(resume_bytes) > MAX_FILE_SIZE_BYTES:
            emit('error', {'message': f'File size exceeds the limit of {MAX_FILE_SIZE_BYTES / (1024 * 1024):.0f} MB.'}, room=request.sid)
            return

        resume_text = None
        if file_type == 'application/pdf':
            resume_text = extract_text_from_pdf(resume_bytes)
        elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            resume_text = extract_text_from_docx(resume_bytes)
        elif file_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
            resume_text = extract_text_from_xlsx(resume_bytes)
        else:
            emit('error', {'message': 'Unsupported file type. Please upload PDF, DOCX, or XLSX.'}, room=request.sid)
            return

        if not resume_text:
            emit('error', {'message': f'Could not extract text from the provided {file_type} file. It might be empty or corrupted.'}, room=request.sid)
            return

        analysis_prompt = f"""
As an expert resume analyzer, compare the following resume to the provided job description.
Your task is to:
1.  Assign a compatibility score between 1 and 10 (10 being a perfect match).
2.  Identify key areas where the resume could be improved to better align with the job description. Provide specific, actionable suggestions.
3.  Suggest a brief (2-3 sentences) professional summary/objective for the resume, tailored to the job description.
4.  Rewrite the ENTIRE provided resume content to optimize it for the job description and improve the compatibility score. Focus on incorporating relevant keywords, emphasizing achievements pertinent to the role, and making the resume more impactful for this specific job. Maintain the original structure and content where possible, but rephrase and enhance sections as needed. Ensure the revised resume content is coherent and professional.

**IMPORTANT:** Your entire response MUST be a SINGLE JSON object. Do NOT include any other text, explanations, or markdown formatting.
The JSON object MUST have the following keys:
-   `score`: An integer from 1 to 10.
-   `suggestions`: An array of strings, each string being an actionable suggestion.
-   `revised_summary`: A brief (2-3 sentences) suggested professional summary/objective for the resume, tailored to the job description.
-   `ai_revised_full_resume_text`: A string containing the entire rewritten resume content, optimized for the job description.
If you cannot generate valid JSON, provide the literal string '{JSON_ERROR_RESPONSE_LITERAL}'.

Resume:
```
{resume_text}
```

Job Description:
```
{job_description}
```
"""
        analysis_schema = {
            "type": "OBJECT",
            "properties": {
                "score": {"type": "INTEGER", "description": "Compatibility score (1-10)"},
                "suggestions": {"type": "ARRAY", "items": {"type": "STRING"}},
                "revised_summary": {"type": "STRING", "description": "Suggested professional summary/objective"},
                "ai_revised_full_resume_text": {"type": "STRING", "description": "The entire rewritten resume content"}
            },
            "required": ["score", "suggestions", "revised_summary", "ai_revised_full_resume_text"]
        }

        # Submit the LLM call to the eventlet background task directly
        sio.start_background_task(async_lm_studio_call, sid=request.sid, prompt=analysis_prompt, schema=analysis_schema, event_name='analysis_result', original_resume_text=resume_text)

    except Exception as e:
        print(f"ERROR: Exception during upload_resume_and_jd: {e}")
        traceback.print_exc()
        emit('error', {'message': f'Server error during analysis: {e}'}, room=request.sid)


@sio.on('apply_edits')
def handle_apply_edits(data):
    """
    Receives edited resume text (now AI-generated and potentially user-tweaked),
    revised summary, and template choice, then generates a new PDF.
    """
    print(f"Received apply_edits event from {request.sid}")
    edited_resume_text = data.get('edited_resume_text')
    revised_summary = data.get('revised_summary')
    selected_template = data.get('selected_template')

    if not edited_resume_text:
        emit('error', {'message': 'Edited resume text is empty.'}, room=request.sid)
        return
    if not selected_template:
        emit('error', {'message': 'Please select a resume template.'}, room=request.sid)
        return

    try:
        new_pdf_bytes = apply_edits_to_template(edited_resume_text, revised_summary, selected_template)
        new_pdf_b64 = base64.b64encode(new_pdf_bytes).decode('utf-8')

        emit('updated_resume_pdf', {'pdf_b64': new_pdf_b64}, room=request.sid)

    except Exception as e:
        print(f"ERROR: Exception during apply_edits: {e}")
        traceback.print_exc()
        emit('error', {'message': f'Server error during PDF generation: {e}'}, room=request.sid)

@sio.on('request_interview_prep')
def handle_request_interview_prep(data):
    """
    Generates interview preparation materials based on job description.
    """
    print(f"Received request_interview_prep event from {request.sid}")
    job_description = data.get('job_description')

    if not job_description:
        emit('error', {'message': 'Job Description is required for interview prep.'}, room=request.sid)
        return

    try:
        interview_prompt = f"""
As an expert interview coach, based on the following job description, generate:
1.  5-7 common interview questions that an applicant for this role might face.
2.  For each question, provide a brief, high-level outline of an ideal answer, highlighting what aspects the interviewer is looking for.
3.  3 general interview tips specifically relevant to preparing for an interview for this type of role.

**IMPORTANT:** Your entire response MUST be a SINGLE JSON object. Do NOT include any other text, explanations, or markdown formatting.
The JSON object MUST have the following keys:
-   `interview_questions`: An array of objects, each with `question` (string) and `ideal_answer_outline` (string).
-   `interview_tips`: An array of strings, each string being a general interview tip.
If you cannot generate valid JSON, provide the literal string '{JSON_ERROR_RESPONSE_LITERAL}'.

Job Description:
```
{job_description}
```
"""
        interview_schema = {
            "type": "OBJECT",
            "properties": {
                "interview_questions": {
                    "type": "ARRAY",
                    "items": {
                        "type": "OBJECT",
                        "properties": {
                            "question": {"type": "STRING"},
                            "ideal_answer_outline": {"type": "STRING"}
                        },
                        "required": ["question", "ideal_answer_outline"]
                    }
                },
                "interview_tips": {"type": "ARRAY", "items": {"type": "STRING"}}
            },
            "required": ["interview_questions", "interview_tips"]
        }

        # Submit the LLM call to the eventlet background task directly
        sio.start_background_task(async_lm_studio_call, sid=request.sid, prompt=interview_prompt, schema=interview_schema, event_name='interview_prep_materials')

    except Exception as e:
        print(f"ERROR: Exception during request_interview_prep: {e}")
        traceback.print_exc()
        emit('error', {'message': f'Server error during interview prep: {e}'}, room=request.sid)

@sio.on('start_interview_simulation')
def handle_start_interview_simulation(data):
    """
    Generates a personalized behavioral interview question based on resume and job description.
    """
    print(f"Received start_interview_simulation event from {request.sid}")
    resume_text = data.get('resume_text')
    job_description = data.get('job_description')

    if not resume_text or not job_description:
        emit('error', {'message': 'Resume text and Job Description are required for interview simulation.'}, room=request.sid)
        return

    try:
        simulation_prompt = f"""
As an expert interviewer. Based on the following resume and job description, generate ONE challenging behavioral interview question.
The question should focus on a skill or experience that is crucial for the role but might be a subtle gap or a strong point to elaborate on from the resume.
Frame it as a STAR method question (e.g., "Tell me about a time when...").
Also, provide a brief (1-2 sentences) context or scenario related to the question.

**IMPORTANT:** Your entire response MUST be a SINGLE JSON object. Do NOT include any other text, explanations, or markdown formatting.
The JSON object MUST have
-   `question`: The behavioral interview question (string).
-   `scenario`: A brief contextual scenario for the question (string).
If you cannot generate valid JSON, provide the literal string '{JSON_ERROR_RESPONSE_LITERAL}'.

Resume:
```
{resume_text}
```

Job Description:
```
{job_description}
```
"""
        simulation_schema = {
            "type": "OBJECT",
            "properties": {
                "question": {"type": "STRING"},
                "scenario": {"type": "STRING"}
            },
            "required": ["question", "scenario"]
        }

        # Submit the LLM call to the eventlet background task directly
        sio.start_background_task(async_lm_studio_call, sid=request.sid, prompt=simulation_prompt, schema=simulation_schema, event_name='interview_question')

    except Exception as e:
        print(f"ERROR: Exception during start_interview_simulation: {e}")
        traceback.print_exc()
        emit('error', {'message': f'Server error during simulation: {e}'}, room=request.sid)

@sio.on('get_interview_feedback')
def handle_get_interview_feedback(data):
    """
    Provides real-time feedback on a user's interview answer.
    """
    print(f"Received get_interview_feedback event from {request.sid}")
    question = data.get('question')
    user_answer = data.get('user_answer')
    job_description = data.get('job_description')
    resume_text = data.get('resume_text')

    if not question or not user_answer or not job_description or not resume_text:
        emit('error', {'message': 'Question, user answer, job description, and resume are required for feedback.'}, room=request.sid)
        return

    try:
        feedback_prompt = f"""
As an expert interview coach. Analyze the following user's interview answer for the given question, considering the job description and their resume.
Provide constructive feedback focusing on:
1.  **STAR Method Adherence:** Does the answer follow the Situation, Task, Action, Result structure? Point out missing components.
2.  **Relevance & Impact:** How well does the answer relate to the job description and demonstrate quantifiable impact? Suggest ways to improve relevance and add metrics.
3.  **Clarity & Conciseness:** Is the answer clear, easy to understand, and to the point?
4.  **Overall Suggestion:** A final overall tip for improving the answer.

**IMPORTANT:** Your entire response MUST be a SINGLE JSON object. Do NOT include any other text, explanations, or markdown formatting.
The JSON object MUST have the following keys:
-   `star_feedback`: string
-   `relevance_impact_feedback`: string
-   `clarity_feedback`: string
-   `overall_suggestion`: string
If you cannot generate valid JSON, provide the literal string '{JSON_ERROR_RESPONSE_LITERAL}'.

---
Job Description:
```
{job_description}
```

Resume:
```
{resume_text}
```

Interview Question:
```
{question}
```

User's Answer:
```
{user_answer}
```
---
"""
        feedback_schema = {
            "type": "OBJECT",
            "properties": {
                "star_feedback": {"type": "STRING"},
                "relevance_impact_feedback": {"type": "STRING"},
                "clarity_feedback": {"type": "STRING"},
                "overall_suggestion": {"type": "STRING"}
            },
            "required": ["star_feedback", "relevance_impact_feedback", "clarity_feedback", "overall_suggestion"]
        }

        # Submit the LLM call to the eventlet background task directly
        sio.start_background_task(async_lm_studio_call, sid=request.sid, prompt=feedback_prompt, schema=feedback_schema, event_name='interview_feedback')

    except Exception as e:
        print(f"ERROR: Exception during get_interview_feedback: {e}")
        traceback.print_exc()
        emit('error', {'message': f'Server error during feedback: {e}'}, room=request.sid)

@sio.on('request_career_roadmap')
def handle_request_career_roadmap(data):
    """
    Generates a predictive career trajectory and skill pathing roadmap.
    """
    print(f"Received request_career_roadmap event from {request.sid}")
    resume_text = data.get('resume_text')
    job_description = data.get('job_description')
    desired_roles = data.get('desired_roles')

    if not resume_text or not job_description or not desired_roles:
        emit('error', {'message': 'Resume text, current job description, and desired roles are required for career roadmap.'}, room=request.sid)
        return

    desired_roles_str = ", ".join(desired_roles)

    try:
        roadmap_prompt = f"""
As an expert career coach and talent development specialist.
Based on the provided resume and the current target job description, and the user's desired future career roles ({desired_roles_str}), generate a personalized career roadmap.

Your roadmap should:
1.  **Identify Key Skill Gaps:** What specific skills or experiences are missing from the resume for the user to progress towards "{desired_roles_str}"? Categorize these (e.g., Technical, Leadership, Communication, etc.).
2.  **Suggest Targeted Learning:** For each identified skill gap, recommend 1-2 specific types of online courses, certifications, or learning resources (e.g., "Advanced Python course on Coursera," "PMP Certification," "Public Speaking workshop").
3.  **Recommend Practical Projects/Experiences:** Suggest specific types of projects or work experiences (personal or professional) that would help bridge these gaps and make the resume more compelling for the desired roles.
4.  **Networking/Soft Skill Advice:** Provide 2-3 actionable tips related to networking, mentorship, or developing crucial soft skills for career advancement.
5.  **Aspirational Summary/Next Steps:** A brief summary of how their resume could look after implementing these steps, and a concluding encouraging remark.

**IMPORTANT:** Your entire response MUST be a SINGLE JSON object. Do NOT include any other text, explanations, or markdown formatting.
The JSON object MUST have the following keys:
-   `skill_gaps`: An array of objects, each with `category` (string) and `skills` (array of strings).
-   `learning_recommendations`: An array of objects, each with `skill_area` (string) and `resources` (array of strings, e.g., "Course: AWS Certified Solutions Architect - Associate (Coursera)").
-   `project_ideas`: An array of strings, each being a project idea.
-   `networking_tips`: An array of strings, each being a networking/soft skill tip.
-   `aspirational_summary`: string
If you cannot generate valid JSON, provide the literal string '{JSON_ERROR_RESPONSE_LITERAL}'.

---
Resume:
```
{resume_text}
```

Current Job Description:
```
{job_description}
```

Desired Future Roles:
```
{desired_roles_str}
```
---
"""
        roadmap_schema = {
            "type": "OBJECT",
            "properties": {
                "skill_gaps": {
                    "type": "ARRAY",
                    "items": {
                        "type": "OBJECT",
                        "properties": {
                            "category": {"type": "STRING"},
                            "skills": {"type": "ARRAY", "items": {"type": "STRING"}}
                        },
                        "required": ["category", "skills"]
                    }
                },
                "learning_recommendations": {
                    "type": "ARRAY",
                    "items": {
                        "type": "OBJECT",
                        "properties": {
                            "skill_area": {"type": "STRING"},
                            "resources": {"type": "ARRAY", "items": {"type": "STRING"}}
                        },
                        "required": ["skill_area", "resources"]
                    }
                },
                "project_ideas": {"type": "ARRAY", "items": {"type": "STRING"}},
                "networking_tips": {"type": "ARRAY", "items": {"type": "STRING"}},
                "aspirational_summary": {"type": "STRING"}
            },
            "required": ["skill_gaps", "learning_recommendations", "project_ideas", "networking_tips", "aspirational_summary"]
        }

        # Submit the LLM call to the eventlet background task directly
        sio.start_background_task(async_lm_studio_call, sid=request.sid, prompt=roadmap_prompt, schema=roadmap_schema, event_name='career_roadmap')

    except Exception as e:
        print(f"ERROR: Exception during request_career_roadmap: {e}")
        traceback.print_exc()
        emit('error', {'message': f'Server error during career roadmap: {e}'}, room=request.sid)


# --- Main entry point ---
if __name__ == '__main__':
    # When using Flask-SocketIO with async_mode='eventlet', you run the app via sio.run()
    # If running with a production WSGI server like Gunicorn, you would configure Gunicorn
    # to use the eventlet/gevent worker class and bind it to Flask-SocketIO's app.
    # Example for Gunicorn: gunicorn --worker-class eventlet -w 1 main:app
    print("Starting Flask-SocketIO server...")
    try:
        sio.run(app, host='0.0.0.0', port=5000, debug=True)
    finally:
        # The ThreadPoolExecutor is no longer directly managed as eventlet.monkey_patch()
        # handles cooperative multitasking for I/O. If there were other CPU-bound tasks
        # not related to I/O that needed true threading, the executor would be useful.
        # For now, it's removed as it's not needed for this particular issue.
        pass
